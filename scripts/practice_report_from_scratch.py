#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from __future__ import annotations

import shutil
import subprocess
from datetime import date, datetime
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parent.parent
DOCS = ROOT / "docs"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(13)


def add_paragraphs(doc: Document, *paragraphs: str) -> None:
    for text in paragraphs:
        add_paragraph(doc, text)


def add_subheading_paragraph(doc: Document, text: str) -> None:
    """Подзаголовок без уровня в многоуровневой нумерации Word (не даёт 2.2.1–2.2.7)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)


def add_code_listing(
    doc: Document,
    rel_path: str,
    *,
    max_lines: int = 260,
    intro: str | None = None,
) -> None:
    """Вставка реального текста файла моноширинным шрифтом (листинг для отчёта)."""
    path = (ROOT / rel_path.replace("\\", "/")).resolve()
    if not path.is_file():
        add_paragraph(doc, f"[Файл не найден в проекте: {rel_path}]")
        return
    try:
        raw = path.read_text(encoding="utf-8")
    except OSError as exc:
        add_paragraph(doc, f"[Не удалось прочитать {rel_path}: {exc}]")
        return
    lines = raw.splitlines()
    total = len(lines)
    if len(lines) > max_lines:
        lines = lines[:max_lines] + [
            "",
            f"— Обрезано для отчёта: в файле всего {total} строк. Полный код: {rel_path} —",
        ]
    block = "\n".join(lines)
    add_paragraph(doc, intro or f"Листинг — исходный код из файла {rel_path}:")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(block)
    run.font.name = "Courier New"
    run.font.size = Pt(8)


def add_photo_here(doc: Document, caption: str, what_on_screen: str) -> None:
    """После листинга — одна строка: вставить фото в Word и как подписать."""
    add_paragraph(doc, f"Вставить скриншот. Подпись под рисунком: «{caption}». На снимке: {what_on_screen}")


def add_section_2_development(doc: Document) -> None:
    """Раздел 2 — подробно по ТЗ и проекту POTOK до раздела 3."""
    add_heading(doc, "2. Разработка веб-приложения (или его части)", level=1)
    add_paragraphs(
        doc,
        "После формулировки целей и задач практики практическая работа была выстроена от документа к результату: "
        "исходной точкой стало оформленное техническое задание (ТЗ): docs/tehnicheskoe_zadanie_potok.pdf. В ТЗ зафиксированы назначение платформы, "
        "границы MVP, роли пользователей, функциональные требования к видео, категориям, комментариям, уведомлениям "
        "и real-time, а также нефункциональные требования к модульности и разделению клиентской и серверной логики.",
        "На основе этого ТЗ был разработан макет интерфейса в Figma: определены каркас приложения (шапка, боковая "
        "навигация, контентная область), ключевые экраны главной ленты, поиска, просмотра видео, страницы канала и "
        "студии автора. Макет не претендовал на финальный брендинг — в соответствии с ограничениями ТЗ акцент сделан "
        "на понятной структуре сценариев и на том, как пользователь переходит от ленты к просмотру и от просмотра к "
        "взаимодействию с каналом и комментариями.",
        "Далее макет был перенесён в кодовую базу Next.js (App Router): общий каркас оформлен в корневом макете "
        "web/src/app/layout.tsx, а функциональные области разбиты на страницы маршрутов и переиспользуемые компоненты "
        "в каталогах web/src/components/* по доменам (layout, home, watch, search, studio и др.). Ниже раздел 2 "
        "раскрывается по подпунктам: сначала содержание технического задания как документа, затем пошагово — "
        "графический интерфейс, база данных, авторизация и работа с серверной частью и данными.",
    )

    add_heading(doc, "2.1 Техническое задание на разработку веб-приложения", level=2)
    add_subheading_paragraph(doc, "Общие сведения")
    add_paragraphs(
        doc,
        "Полное наименование разрабатываемого программного продукта в рамках практики: веб-приложение POTOK — "
        "платформа публикации и просмотра видео с элементами социального взаимодействия (каналы, подписки, комментарии, "
        "уведомления). Нормативный документ — техническое задание (ТЗ) в оформленном виде: docs/tehnicheskoe_zadanie_potok.pdf "
        "(печать и приложение к отчёту).",
        "Плановые сроки начала и окончания работ по разработке в рамках производственной практики определяются "
        "графиком учебного заведения и должны укладываться в период прохождения практики (конкретные даты указываются "
        "в приказе и дневнике практики исполнителем вручную).",
    )

    add_subheading_paragraph(doc, "Назначение, цели и задачи создания сайта (в связке с ОК и ПК)")
    add_paragraphs(
        doc,
        "Назначение POTOK в редакции ТЗ — предоставить пользователю доступ к каталогу видео, просмотру записи, поиску и "
        "социальным действиям (комментарии, реакции, подписки), а автору контента — средства публикации метаданных и "
        "медиа через облачное хранилище, управления каналом и обратной связи через уведомления. В проекте автоматизируется "
        "сбор и отображение метаданных о роликах, учёт просмотров и взаимодействий, доставка событий в реальном времени "
        "для комментариев и уведомлений.",
        "Задачи создания сайта сформулированы с опорой на общие и профессиональные компетенции: необходимо было "
        "проанализировать предметную область и требования (ОК), спроектировать структуру данных и интерфейсные сценарии "
        "(ПК), реализовать клиентскую часть на React/Next.js с типизацией TypeScript (ПК), настроить взаимодействие с "
        "PostgreSQL через Supabase с политиками RLS (ПК), обеспечить аутентификацию и разграничение действий пользователя "
        "(ПК), внедрить канал real-time через Pusher согласно ТЗ (ПК), проверить ключевые пользовательские сценарии и "
        "зафиксировать решения в отчёте (ОК).",
        "Проект призван решать задачи «цифровой витрины» видеоконтента и учебного стенда для отработки полного цикла "
        "веб-разработки: от постановки задачи до хранения данных и отображения их в интерфейсе.",
    )

    add_subheading_paragraph(doc, "Функциональные подсистемы (по смыслу ТЗ)")
    add_paragraphs(
        doc,
        "В логике ТЗ и реализации POTOK можно выделить три взаимосвязанные подсистемы. Подсистема обработки данных "
        "охватывает формы загрузки видео, ввод метаданных, действия пользователя (лайк, подписка, комментарий) и "
        "серверные обработчики, которые принимают эти действия и приводят к изменению состояния в базе. Подсистема "
        "хранения данных реализована в Supabase PostgreSQL: таблицы пользователей и профилей, видео, категорий, "
        "подписок, комментариев, уведомлений и связанные ограничения и индексы, описанные в SQL-скриптах каталога "
        "web/supabase. Подсистема формирования и визуализации включает страницы ленты, поиска, просмотра, канала, "
        "студии и виджеты уведомлений: она отображает агрегированную информацию из БД и медиа-ссылки из Storage.",
    )

    add_subheading_paragraph(doc, "Состав и содержание работ по созданию сайта (обобщённо по этапам ТЗ)")
    add_paragraphs(
        doc,
        "Работы по созданию сайта в соответствии с разделом «Этапы реализации» технического задания (ТЗ) можно сгруппировать в три "
        "крупных блока, которые далее раскрываются в подразделах 2.2–2.5.",
        "Первый блок — инфраструктура и данные: инициализация проекта Next.js, переменные окружения для Supabase и "
        "Pusher, создание схемы БД, индексов и RLS-политик. Второй блок — пользователь и контент: аутентификация "
        "Supabase Auth, профиль и канал, загрузка файлов в Storage и запись метаданных в таблицу videos, страница "
        "просмотра с плеером Plyr. Третий блок — социальные функции и качество: категории и лента, подписки и реакции, "
        "комментарии с real-time, уведомления, пагинация и обработка ошибок. Конкретные сроки подэтапов определяются "
        "планом работ и договором на практику; в отчёте фиксируется фактическая последовательность реализации в коде.",
    )

    add_heading(doc, "2.2 Разработка графического интерфейса веб-приложения", level=2)
    add_paragraphs(
        doc,
        "Графический интерфейс развивался от макета в Figma к компонентной вёрстке на React: сначала зафиксированы "
        "повторяемые блоки (шапка, сайдбар, карточка ролика, область плеера, формы), затем они собраны в страницы App Router. "
        "При вёрстке учитывались требования ТЗ к модульности: разделение layout и контента, клиентские компоненты там, "
        "где нужны хуки и подписки, серверные страницы там, где возможна предварительная загрузка данных.",
        "Визуальный стиль: выбрана тёмная палитра под видеоконтент (глубокий фон панелей и контентной области, светлый основной текст и приглушённые подписи второго уровня). Акценты и интерактивные границы выделены холодными оттенками cyan и синего, чтобы навигация читалась на тёмном фоне при длительном просмотре. Шрифты задаются через Next.js (Geist Sans / моноширинный для кода), размеры и отступы — утилитами Tailwind CSS; контраст текста и фона проверялся визуально на типичных экранах. Для кнопок, ссылок и полей ввода предусмотрены состояния наведения и фокуса. Адаптивность обеспечивается брейкпоинтами Tailwind: сетки колонок меняют число рядов по ширине, навигация переключается в мобильный режим (выезжающая панель, кнопка меню в шапке). В глобальных стилях учтена системная тёмная тема (`prefers-color-scheme: dark`) и безопасные отступы под экраны с вырезом (`safe-area`).",
        "Ниже каждый подраздел оформлен так: сначала пояснение, что именно реализовано и зачем; затем вставлен реальный "
        "листинг из файлов проекта (не путь «на бумаге», а текст кода); сразу после листинга — иллюстрация того же экрана. "
        "Так у проверяющего в одном месте и код, и визуальный результат.",
        "Важно: код самого веб-приложения (TypeScript, React, Next.js) сосредоточен в разделе 2.2 и в фрагментах 2.4–2.5; "
        "раздел 2.3 содержит только SQL и политики Supabase — это слой базы данных, а не интерфейса.",
    )

    add_subheading_paragraph(doc, "Ядро клиентского приложения: стек, конфигурация, оболочка и авторизация")
    add_paragraphs(
        doc,
        "Перед разбором отдельных экранов приведены файлы, которые определяют стек и «скелет» сайта: зависимости npm, "
        "настройки Next.js (маршруты, rewrites), компонент AppShell, который собирает сайдбар и контент, и контекст "
        "аутентификации для всего клиента. Без этого остальные экраны раздела 2.2 ниже не работают как единое приложение.",
    )
    add_code_listing(
        doc,
        "web/package.json",
        max_lines=40,
        intro="Листинг 2.2а — web/package.json (зависимости: Next.js, React, Supabase, Pusher, Plyr, Zod):",
    )
    add_code_listing(doc, "web/next.config.ts", intro="Листинг 2.2б — web/next.config.ts (конфигурация Next.js, rewrites):")
    add_code_listing(
        doc,
        "web/src/components/layout/app-shell.tsx",
        intro="Листинг 2.2в — web/src/components/layout/app-shell.tsx (оболочка: сайдбар + область страницы):",
    )
    add_code_listing(
        doc,
        "web/src/components/auth/auth-context.tsx",
        max_lines=120,
        intro="Листинг 2.2г — web/src/components/auth/auth-context.tsx (фрагмент: состояние сессии Supabase на клиенте):",
    )

    add_subheading_paragraph(doc, "Каркас приложения, шапка и боковая навигация")
    add_paragraphs(
        doc,
        "Корневой layout подключает глобальные стили и оболочку AppShell: так на всех страницах одинаковая высота, шрифты "
        "и поведение боковой панели. Шапка AppHeader содержит бренд, поиск и действия пользователя; сайдбар даёт "
        "постоянные разделы. Ниже приведён полный текст корневого layout, затем фрагмент шапки (клиентский компонент с "
        "запросом к профилю) и полный файл сайдбара.",
    )
    add_code_listing(doc, "web/src/app/layout.tsx", intro="Листинг 2.2а — web/src/app/layout.tsx (корневой макет):")
    add_code_listing(
        doc,
        "web/src/components/layout/app-header.tsx",
        max_lines=90,
        intro="Листинг 2.2б — web/src/components/layout/app-header.tsx (начало файла, шапка и логика профиля):",
    )
    add_code_listing(doc, "web/src/components/layout/sidebar.tsx", intro="Листинг 2.2в — web/src/components/layout/sidebar.tsx:")
    add_photo_here(
        doc,
        "Рисунок 2.2 — шапка и боковая навигация POTOK",
        "браузер: шапка с ПОТОК и поиском, слева открытый сайдбар с пунктами меню.",
    )

    add_subheading_paragraph(doc, "Главная страница: лента и категории")
    add_paragraphs(
        doc,
        "Главная страница соответствует сценарию ТЗ: пользователь попадает в ленту видео; при необходимости маршрут "
        "перенаправляет на поиск или историю. Страница собирает AppHeader и HomeVideoFeed — отдельный компонент, "
        "чтобы вынести логику выборки и отображения карточек.",
    )
    add_code_listing(doc, "web/src/app/page.tsx", intro="Листинг 2.2а — web/src/app/page.tsx:")
    add_code_listing(
        doc,
        "web/src/components/home/home-video-feed.tsx",
        intro="Листинг 2.2б — web/src/components/home/home-video-feed.tsx:",
    )
    add_photo_here(
        doc,
        "Рисунок 2.2 — главная страница POTOK",
        "главная /: шапка и лента карточек видео.",
    )

    add_subheading_paragraph(doc, "Поиск по контенту")
    add_paragraphs(
        doc,
        "Поиск вынесен в отдельный маршрут; результаты выводятся отдельным компонентом, чтобы позже добавить пагинацию и "
        "фильтры без переписывания страницы.",
    )
    add_code_listing(doc, "web/src/app/search/page.tsx", intro="Листинг 2.2а — web/src/app/search/page.tsx:")
    add_code_listing(
        doc,
        "web/src/components/search/search-results.tsx",
        max_lines=110,
        intro="Листинг 2.2б — web/src/components/search/search-results.tsx (фрагмент):",
    )
    add_photo_here(
        doc,
        "Рисунок 2.2 — страница поиска",
        "/search: строка запроса и результаты или пустая выдача.",
    )

    add_subheading_paragraph(doc, "Страница просмотра видео, плеер и комментарии")
    add_paragraphs(
        doc,
        "Страница просмотра связывает динамический маршрут watch/[id], плеер на Plyr и блок комментариев. Это ключевой "
        "сценарий ТЗ: просмотр записи и обсуждение под ней. Ниже приведены фрагменты кода страницы, плеера и комментариев, "
        "затем скрин того же экрана.",
    )
    add_code_listing(
        doc,
        "web/src/app/watch/[id]/page.tsx",
        max_lines=120,
        intro="Листинг 2.2а — web/src/app/watch/[id]/page.tsx (фрагмент):",
    )
    add_code_listing(
        doc,
        "web/src/components/watch/watch-player.tsx",
        max_lines=100,
        intro="Листинг 2.2б — web/src/components/watch/watch-player.tsx (фрагмент):",
    )
    add_code_listing(
        doc,
        "web/src/components/watch/comments-section.tsx",
        intro="Листинг 2.2в — web/src/components/watch/comments-section.tsx:",
    )
    add_photo_here(
        doc,
        "Рисунок 2.2 — страница просмотра видео",
        "страница /watch/…: плеер, название, блок комментариев.",
    )

    add_subheading_paragraph(doc, "Страница канала")
    add_paragraphs(
        doc,
        "Страница канала отражает связь аккаунта и публичного представления автора (маршрут с параметром handle). "
        "Ниже — фрагмент кода страницы и скрин интерфейса канала.",
    )
    add_code_listing(
        doc,
        "web/src/app/channel/[handle]/page.tsx",
        max_lines=130,
        intro="Листинг 2.2 — web/src/app/channel/[handle]/page.tsx (фрагмент):",
    )
    add_photo_here(
        doc,
        "Рисунок 2.2 — страница канала автора",
        "канал: шапка (аватар, название) и видео канала.",
    )

    add_subheading_paragraph(doc, "Студия автора (загрузка и управление контентом)")
    add_paragraphs(
        doc,
        "Студия концентрирует сценарии автора: загрузка, управление контентом, дополнительные разделы. Навигация между "
        "вкладками вынесена в studio-sidebar. Приведены фрагменты кода и скрин студии.",
    )
    add_code_listing(
        doc,
        "web/src/app/studio/page.tsx",
        max_lines=130,
        intro="Листинг 2.2а — web/src/app/studio/page.tsx (фрагмент):",
    )
    add_code_listing(
        doc,
        "web/src/components/studio/studio-sidebar.tsx",
        max_lines=120,
        intro="Листинг 2.2б — web/src/components/studio/studio-sidebar.tsx (фрагмент):",
    )
    add_photo_here(
        doc,
        "Рисунок 2.2 — студия автора",
        "/studio: панель вкладок слева и область контента.",
    )

    add_heading(doc, "2.3 Разработка базы данных веб-приложения", level=2)
    add_subheading_paragraph(doc, "Установка и настройка среды хранения данных (Supabase вместо Laravel/WordPress)")
    add_paragraphs(
        doc,
        "В отличие от учебных примеров на Laravel или WordPress, в POTOK используется облачный стек Supabase: "
        "управляемый PostgreSQL, встроенная аутентификация и API. В консоли Supabase создаётся проект, копируются URL "
        "и ключи anon и service role, настраивается Storage для видео и превью; в приложении значения задаются в "
        ".env.local. Подключение из браузера и с сервера выполняется через @supabase/ssr и createClient.",
        "Ниже — реальный код клиентов Supabase (браузер и сервер). После листингов — место для скриншота из консоли.",
    )
    add_code_listing(doc, "web/src/lib/supabase/client.ts", intro="Листинг 2.3а — web/src/lib/supabase/client.ts:")
    add_code_listing(doc, "web/src/lib/supabase/server.ts", intro="Листинг 2.3б — web/src/lib/supabase/server.ts:")
    add_photo_here(
        doc,
        "Рисунок 2.3 — консоль Supabase",
        "supabase.com: проект — настройки API (ключи можно замазать) или раздел Storage.",
    )

    add_subheading_paragraph(doc, "Создание базы данных: таблицы, связи, миграции")
    add_paragraphs(
        doc,
        "Структура данных следует техническому заданию (ТЗ, docs/tehnicheskoe_zadanie_potok.pdf): профили пользователей, категории, видео, подписки, лайки, комментарии, "
        "уведомления. Дополнительно в проекте добавлены сущности для плейлистов и триггеры целостности — см. отдельные "
        "SQL-скрипты в каталоге web/supabase.",
        "Логическая схема связей (словесно, для отчёта): таблица public.users ссылается на auth.users и один-к-одному "
        "представляет канал; public.videos ссылается на users и categories; comments ссылаются на videos и parent-комментарий; "
        "subscriptions связывают подписчика и канал с ограничением «не на себя»; likes — по паре пользователь–видео; "
        "notifications хранятся на пользователя. Такая модель отражена в SQL ниже.",
        "Далее — три полных листинга SQL из проекта, затем место для скриншота из Supabase.",
    )
    add_code_listing(doc, "web/supabase/01_core.sql", intro="Листинг 2.3а — web/supabase/01_core.sql (ядро таблиц и справочник категорий):")
    add_code_listing(doc, "web/supabase/03_rls_policies.sql", intro="Листинг 2.3б — web/supabase/03_rls_policies.sql (политики RLS):")
    add_code_listing(
        doc,
        "web/supabase/12_playlists_system_and_triggers.sql",
        intro="Листинг 2.3в — web/supabase/12_playlists_system_and_triggers.sql (триггеры и плейлисты):",
    )
    add_photo_here(
        doc,
        "Рисунок 2.3 — таблицы или SQL в Supabase",
        "Table editor (список таблиц / колонки) или SQL Editor с успешным запросом.",
    )

    add_subheading_paragraph(doc, "Маршруты, запросы и серверная логика (аналог контроллеров в Next.js)")
    add_paragraphs(
        doc,
        "В App Router роль контроллеров выполняют серверные route handlers и серверные компоненты страниц. Запросы к "
        "данным идут через Supabase SDK; для внешних событий используется Pusher. Ниже приведён полный код серверного "
        "маршрута, который триггерит событие в канале Pusher.",
    )
    add_code_listing(
        doc,
        "web/src/app/api/realtime/pusher/trigger/route.ts",
        intro="Листинг 2.3 — web/src/app/api/realtime/pusher/trigger/route.ts:",
    )

    add_subheading_paragraph(doc, "Валидация данных на клиенте")
    add_paragraphs(
        doc,
        "Формы авторизации и ввод данных на клиенте сопровождаются проверкой через Zod (см. auth и API-маршруты): "
        "неверно заполненные поля отсекаются до отправки на сервер. Ниже приведён фрагмент кода страницы входа, где "
        "заданы схемы и обработка ошибок.",
    )
    add_code_listing(
        doc,
        "web/src/app/auth/page.tsx",
        max_lines=130,
        intro="Листинг 2.3 — web/src/app/auth/page.tsx (фрагмент: импорты, схемы Zod, начало форм):",
    )
    add_photo_here(
        doc,
        "Рисунок 2.3 — ошибка валидации (по желанию)",
        "/auth: красная ошибка под полем после неверного ввода.",
    )

    add_heading(doc, "2.4 Разработка механизмов авторизации и аутентификации пользователей", level=2)
    add_paragraphs(
        doc,
        "Аутентификация реализована через Supabase Auth: регистрация и вход по email и паролю. Пароли хранит и хэширует "
        "платформа Auth, не прикладная таблица. Восстановление доступа — сценарий «email → код из письма → новый пароль» "
        "на страницах forgot-password и reset-password. Шаблон письма для Supabase лежит в репозитории.",
        "Ниже — листинги кода; после них — места для скриншотов экранов входа и восстановления пароля.",
    )
    add_subheading_paragraph(doc, "Вход, регистрация и восстановление пароля (код и иллюстрации)")
    add_code_listing(
        doc,
        "web/src/app/auth/page.tsx",
        max_lines=100,
        intro="Листинг 2.4а — web/src/app/auth/page.tsx (фрагмент начала файла):",
    )
    add_code_listing(
        doc,
        "web/src/app/auth/forgot-password/page.tsx",
        max_lines=130,
        intro="Листинг 2.4б — web/src/app/auth/forgot-password/page.tsx (фрагмент):",
    )
    add_code_listing(
        doc,
        "web/src/app/auth/reset-password/page.tsx",
        max_lines=100,
        intro="Листинг 2.4в — web/src/app/auth/reset-password/page.tsx (фрагмент):",
    )
    add_code_listing(
        doc,
        "web/supabase/email-templates/recovery.html",
        intro="Листинг 2.4г — web/supabase/email-templates/recovery.html (шаблон письма для консоли Supabase):",
    )
    add_photo_here(
        doc,
        "Рисунок 2.4 — вход и регистрация",
        "/auth: форма входа и регистрации.",
    )
    add_photo_here(
        doc,
        "Рисунок 2.4 — код из письма",
        "/auth/forgot-password: экран ввода кода.",
    )
    add_photo_here(
        doc,
        "Рисунок 2.4 — письмо с кодом (по желанию)",
        "почта: письмо от Supabase с кодом.",
    )

    add_heading(doc, "2.5 Работа с серверной частью программного продукта. Отправка и получение данных из БД", level=2)
    add_paragraphs(
        doc,
        "Чтение и запись в PostgreSQL выполняются через Supabase: в браузере — с anon-ключом и RLS, на сервере при "
        "необходимости — service role-клиент для доверенных операций. Медиа хранятся в Storage, в таблицах — только "
        "метаданные и URL, как в ТЗ. Real-time для комментариев обеспечивается через Pusher и серверный триггер.",
        "Ниже — листинги сервисного клиента и клиентского вызова Pusher; затем место для скриншота из DevTools.",
    )
    add_subheading_paragraph(doc, "Сервисный клиент, Pusher и маршрут API")
    add_code_listing(doc, "web/src/lib/supabase/service.ts", intro="Листинг 2.5а — web/src/lib/supabase/service.ts:")
    add_code_listing(doc, "web/src/lib/pusher/trigger.ts", intro="Листинг 2.5б — web/src/lib/pusher/trigger.ts (клиентский вызов API):")
    add_paragraphs(
        doc,
        "Серверный обработчик POST для Pusher приведён в листинге 2.3 (тот же файл web/src/app/api/realtime/pusher/trigger/route.ts) — "
        "здесь повторно не дублируется, чтобы не раздувать отчёт.",
    )
    add_photo_here(
        doc,
        "Рисунок 2.5 — запрос к API в браузере",
        "F12 → Сеть: POST /api/realtime/pusher/trigger, статус 200.",
    )


def build_docx(out_docx: Path) -> None:
    DOCS.mkdir(parents=True, exist_ok=True)
    today = date.today().strftime("%d.%m.%Y")

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(13)

    add_heading(doc, "ОТЧЁТ О ПРОХОЖДЕНИИ ПРОИЗВОДСТВЕННОЙ ПРАКТИКИ", level=0)
    add_paragraph(doc, f"Проект: POTOK")
    add_paragraph(doc, f"Дата формирования отчёта: {today}")
    add_heading(doc, "Содержание", level=1)
    add_paragraph(
        doc,
        "Оформите оглавление в Word: вкладка «Ссылки» → «Оглавление», либо вставьте сюда скриншот оглавления по образцу "
        "методички (разделы 1, 2, 3 и заключение с номерами страниц).",
    )
    add_heading(doc, "1. Цель, задачи, описание базы прохождения производственной практики", level=1)
    add_heading(doc, "1.1 Цель и задачи производственной практики", level=2)
    add_paragraph(
        doc,
        "Целью производственной практики является применение теоретических знаний, полученных при изучении "
        "профессионального модуля ПМ 09 «Проектирование, разработка и оптимизация веб-приложений», при разработке "
        "реального программного продукта POTOK. Практика позволила перевести учебные знания в прикладной результат, "
        "сформировать устойчивые профессиональные навыки и закрепить компетенции, связанные с анализом, "
        "проектированием, разработкой и сопровождением веб-приложений."
    )
    add_paragraph(
        doc,
        "Основные задачи практики заключались в проектировании интерфейса, реализации ключевых пользовательских "
        "сценариев, интеграции клиентской части с серверными сервисами, обеспечении базовой безопасности и "
        "первичной оптимизации производительности. Дополнительной задачей являлась подготовка отчётной документации "
        "с листингом и пояснением выполненных решений."
    )

    add_heading(doc, "1.2 Общее ознакомление с организацией и базой практики", level=2)
    add_paragraph(
        doc,
        "Практика проходила на базе учебного проекта в репозитории potok в сроки, предусмотренные графиком колледжа. "
        "В качестве объекта разработки использовалась веб-платформа видеоконтента с функционалом просмотра, поиска, "
        "комментариев, канала, студии автора и механизмов авторизации."
    )
    add_paragraph(
        doc,
        "Технологическая база проекта включает Next.js, TypeScript, Supabase, Pusher и Plyr. Организация работы "
        "строилась по инженерной дисциплине: контроль версий через git, разделение конфигурации через env-переменные, "
        "валидация пользовательского ввода и проверка ключевых сценариев после изменений."
    )

    add_section_2_development(doc)

    add_heading(doc, "3. Оптимизация и безопасность веб-приложения", level=1)
    add_heading(doc, "3.1 Внутренняя SEO оптимизация", level=2)
    add_paragraph(
        doc,
        "В рамках практики выполнена базовая SEO-подготовка структуры страниц и маршрутов, что создаёт основу для "
        "дальнейшего роста видимости проекта."
    )
    add_heading(doc, "3.2 Общий аудит: SEO, юзабилити, тексты", level=2)
    add_paragraph(
        doc,
        "Проведён внутренний аудит пользовательских сценариев, текстов форм и логики навигации для повышения "
        "понятности интерфейса."
    )
    add_heading(doc, "3.3 Исследование способов ускорения загрузки", level=2)
    add_paragraph(
        doc,
        "Определены направления оптимизации: контроль размера клиентского бандла, оптимизация медиа-ресурсов, "
        "декомпозиция тяжёлых интерфейсных блоков и профилирование критичных участков."
    )
    add_heading(doc, "3.4 Проверка безопасности", level=2)
    add_paragraph(
        doc,
        "Рассмотрены риски clickjacking, CSRF, XSS и SQL-инъекций, а также отмечена необходимость отдельного прогона "
        "автоматического сканирования средствами ZAP в тестовом контуре."
    )
    add_photo_here(
        doc,
        "Рисунок 3.4 — безопасность (по желанию)",
        "ZAP с отчётом или DevTools → заголовки ответа главной страницы.",
    )

    add_heading(doc, "4. Заключение", level=1)
    add_paragraph(
        doc,
        "В результате практики реализован рабочий MVP веб-платформы POTOK в соответствии с техническим заданием (ТЗ) "
        "— см. docs/tehnicheskoe_zadanie_potok.pdf. "
        "Получен практический опыт полного цикла веб-разработки: от проектирования интерфейса до интеграции с "
        "серверными сервисами и первичной оценки качества решения."
    )
    add_paragraph(doc, "Исполнитель __________________ дата __________")

    try:
        doc.save(out_docx)
    except PermissionError:
        alt = out_docx.with_name(out_docx.stem + "_new" + out_docx.suffix)
        doc.save(alt)
        print(f"WARN: {out_docx.name} занят (Word?). Сохранено как: {alt.name}")


def export_html_and_pdf(out_html: Path, out_pdf: Path, docx_name: str) -> None:
    html = f"""<!doctype html><html lang='ru'><head><meta charset='utf-8'><title>Отчёт POTOK</title>
<style>body{{font-family:'Segoe UI',Arial,sans-serif;max-width:900px;margin:24px auto;line-height:1.5;padding:0 12px;}}</style>
</head><body><h1>Отчёт сгенерирован в DOCX</h1><p>Основной файл: {docx_name}</p></body></html>"""
    out_html.write_text(html, encoding="utf-8")

    browsers = [
        shutil.which("msedge"),
        shutil.which("chrome"),
        r"C:\\Program Files (x86)\\Microsoft\\Edge\\Application\\msedge.exe",
        r"C:\\Program Files\\Microsoft\\Edge\\Application\\msedge.exe",
        r"C:\\Program Files\\Google\\Chrome\\Application\\chrome.exe",
    ]
    browser = next((b for b in browsers if b and Path(b).exists()), None)
    if browser:
        subprocess.run(
            [str(browser), "--headless", "--disable-gpu", f"--print-to-pdf={out_pdf}", str(out_html)],
            capture_output=True,
            text=True,
            check=False,
        )


def main() -> None:
    DOCS.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_docx = DOCS / f"otchet_praktika_{stamp}.docx"
    out_html = DOCS / f"otchet_praktika_{stamp}.html"
    out_pdf = DOCS / f"otchet_praktika_{stamp}.pdf"
    build_docx(out_docx)
    export_html_and_pdf(out_html, out_pdf, out_docx.name)
    latest_docx = DOCS / "otchet_praktika_final.docx"
    try:
        shutil.copy2(out_docx, latest_docx)
    except OSError as e:
        print(f"WARN: не удалось обновить {latest_docx.name}: {e}")
    print(f"DOCX (новый): {out_docx}")
    print(f"PDF (новый):  {out_pdf}")
    print(f"Копия latest: {latest_docx}")


if __name__ == "__main__":
    main()
