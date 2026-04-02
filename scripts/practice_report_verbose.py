#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Подробный отчёт по практике POTOK: сначала развёрнутое повествование о работе,
затем иллюстрация (скриншот), в конце — листинги кода и структуры.

По умолчанию пишет только docs/otchet_praktika_potok.docx. Флаг --all добавляет
.md, .html и при наличии браузера — .pdf.

Положите скриншот раздела отчёта в docs/screenshots/otchet_razdel_2_2.png
(или укажите путь через --screenshot).
"""

from __future__ import annotations

import argparse
import os
import re
import shutil
import sys
from datetime import date
from pathlib import Path

_SCRIPTS_DIR = Path(__file__).resolve().parent
if str(_SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(_SCRIPTS_DIR))

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.image.exceptions import UnrecognizedImageError
from docx.shared import Cm, Inches, Pt

from practice_report import (
    ROOT,
    build_code_listings_plain,
    ensure_docs_dir,
    try_export_pdf_from_html,
)

OUT_MD = ROOT / "docs" / "otchet_praktika_potok.md"
OUT_DOCX = ROOT / "docs" / "otchet_praktika_potok.docx"
OUT_HTML = ROOT / "docs" / "otchet_praktika_potok.html"
OUT_PDF = ROOT / "docs" / "otchet_praktika_potok.pdf"

DEFAULT_SCREENSHOT_REL = "docs/screenshots/otchet_razdel_2_2.png"

IMG_LINE = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$")


def build_listings_plain() -> str:
    """Только листинги кода (без дерева каталогов)."""
    return build_code_listings_plain()


def section_2_2_long_narrative() -> str:
    """Развёрнутый текст по интерфейсу и адаптиву (много связного повествования)."""
    return """
Прежде чем переходить к перечислению файлов и фрагментов исходного кода, целесообразно последовательно и достаточно подробно описать ту работу, которая была выполнена на этапе внедрения пользовательского интерфейса и адаптивной вёрстки, поскольку именно этот слой определяет первое впечатление пользователя и во многом задаёт удобство дальнейшего взаимодействия с платформой POTOK. В рамках практики интерфейс рассматривался не как набор отдельных «картинок», а как целостная система: от корневого макета приложения и глобальных настроек области просмотра до поведения отдельных блоков на экранах разной ширины — от смартфона до широкого монитора.

Исходной точкой оставалось техническое задание и макет в Figma, в котором были зафиксированы основные сценарии: переход от главной ленты к просмотру видео, работа с поиском, навигация по разделам, связанным с каналом и студией автора. Перенос макета в код осуществлялся в экосистеме Next.js с использованием App Router: такой подход позволяет естественно разделить маршруты по каталогу `web/src/app`, а переиспользуемые визуальные и логические блоки — вынести в `web/src/components`. Корневой файл `web/src/app/layout.tsx` задаёт оболочку страницы: подключение глобальных стилей, шрифтов и обёртку `AppShell`, внутри которой уже размещается содержимое конкретных маршрутов. Отдельного внимания заслуживает экспорт `viewport` с параметром `viewportFit: "cover"`: он нужен для того, чтобы на устройствах с вырезом экрана («чёлкой») и скруглёнными углами контент корректно заполнял безопасную область, а системные отступы можно было учесть в компонентах через CSS-переменные окружения `env(safe-area-inset-*)`.

Цвет, типографика и контраст: для платформы POTOK выбрана тёмная палитра под длительный просмотр видео — глубокий фон контентных блоков и панелей, светлый основной текст и приглушённые подписи второго уровня (slate). Акцентные элементы (кнопки, ссылки, обводки карточек, полоса прокрутки в глобальных стилях) выдержаны в холодных оттенках cyan и синего, чтобы визуально отделять интерактив от контента и не перегружать глаз. Шрифты подключаются через Next.js (Geist Sans / моноширинный для кода), базовые размеры и межстрочные интервалы задаются утилитами Tailwind; контраст текста и фона проверялся визуально на типичных экранах. В `globals.css` для `body` заданы переменные фона и текста с учётом `prefers-color-scheme: dark`, чтобы интерфейс согласовывался с системной темой. Для форм и кликабельных элементов предусмотрены состояния наведения и фокуса, что важно и для мыши, и для навигации с клавиатуры.

Каркас приложения, с которым пользователь сталкивается на каждой странице, построен вокруг верхней шапки и боковой панели навигации. Компонент `AppHeader` отвечает за узнаваемый верхний ряд: логотип, центральный блок поиска и действия, связанные с учётной записью. Для узких экранов, когда боковая панель по умолчанию скрыта и не занимает место в потоке, в шапке была добавлена отдельная кнопка с иконкой меню: по нажатию вызывается функция из контекста сайдбара (`useSidebarState`), и пользователь получает привычный паттерн, сходный с мобильным интерфейсом YouTube — навигация не «теряется», а становится доступной по явному запросу. Таким образом, на десктопе сохраняется привычная двухколоночная схема, а на планшете и телефоне интерфейс остаётся читаемым и не перегружен горизонтальными элементами.

Компонент `Sidebar` реализует боковую панель как элемент, который ведёт себя по-разному в зависимости от ширины экрана. На больших ширинах панель участвует в общей сетке страницы, может быть развёрнута (с подписями пунктов) или свёрнута до узкой полосы с иконками. На экранах уже контрольной точки `lg` панель превращается в выезжающий «ящик»: она фиксируется у края, получает собственную прокрутку при длинном списке пунктов и закрывается при выборе маршрута на сенсорном устройстве, чтобы не оставлять полупрозрачный слой поверх контента без необходимости. В стилях заложены отступы с учётом `safe-area` сверху и снизу, чтобы на современных телефонах нижняя часть меню не уезжала под системную полосу жестов.

Главная страница строится как связка полосы категорий и блока рекомендаций. Полоса категорий (`CategoryChipsBar`) визуально напоминает аналог на крупных видеосервисах: горизонтальный ряд чипов с иконками, который можно прокручивать свайпом, а на широких экранах дополнительно подстрахован стрелками прокрутки. Блок закреплён под шапкой с помощью `position: sticky`, причём вертикальное смещение `top` вычисляется с учётом высоты шапки и снова же смещения из safe-area, чтобы при прокрутке ленты категории не «наезжали» на логотип и поиск. Такое решение требует согласованности между высотой шапки и формулой смещения, зато даёт предсказуемое поведение на разных устройствах.

Центральная часть главной — сетка карточек видео (`HomeVideoFeed`). Здесь была настроена ступенчатая адаптация: на самых узких экранах карточки идут в одну колонку, что упрощает нажатие пальцем и чтение заголовков; на промежуточной ширине появляются две колонки; на ноутбуке — три; на очень широких экранах (`xl` и выше) сетка расширяется до четырёх колонок в ряд по аналогии с привычной лентой YouTube, при этом общая ширина контента ограничена разумным максимумом, чтобы на мониторах с большим разрешением строка не растягивалась до нечитаемо редких карточек. Дополнительно нижний отступ контейнера учитывает safe-area снизу, чтобы последний ряд карточек не оказывался под системной навигацией.

Поиск в шапке реализован отдельным компонентом; для корректного сжатия внутри гибкой строки шапки контейнер поиска снабжён классами вроде `min-w-0`, что важно в типичной для Flexbox ситуации, когда без явного ограничения минимальной ширины поле ввода «отказывается» сужаться и выталкивает соседние элементы. Подобные детали кажутся мелкими, но именно они отделяют «просто собранный интерфейс» от интерфейса, который стабильно ведёт себя на реальных устройствах.

После внесения перечисленных изменений выполнялась проверка сборки командой `npm run build` в каталоге веб-приложения. Успешное завершение сборки в режиме production подтверждает отсутствие синтаксических ошибок, согласованность типов TypeScript и корректную работу статического анализа в рамках конфигурации Next.js. Таким образом, блок работ по интерфейсу и адаптиву был не только спроектирован и реализован в коде, но и верифицирован на уровне инструментов сборки, что для отчётной документации по практике является важным практическим аргументом.

Далее по тексту подраздела об интерфейсе приводятся иллюстрации (при наличии файлов) и полный блок листингов исходного кода без перечня структуры папок: SQL и фрагменты TypeScript/React, подтверждающие изложенное.
""".strip()


def build_verbose_report_body(screenshot_rel: str) -> str:
    today = date.today().strftime("%d.%m.%Y")
    listings = build_listings_plain()
    shot_abs = ROOT / screenshot_rel
    shot_md = ""
    if shot_abs.is_file():
        shot_md = (
            f"\n![Иллюстрация к подразделу об интерфейсе]({screenshot_rel})\n\n"
            "*Рисунок.* Фрагмент отчёта или экрана (файл изображения подключён автоматически).\n\n"
        )

    intro_note = """
Этот вариант отчёта — **подробная редакция**: сначала развёрнутое описание работ, затем иллюстрация (если файл есть), затем **только листинги исходного кода** (SQL и приложение), без дерева каталогов.
""".strip()

    return f"""# ОТЧЁТ О ПРОХОЖДЕНИИ ПРОИЗВОДСТВЕННОЙ ПРАКТИКИ (подробная редакция)

Проект POTOK  
Дата формирования отчёта {today}

**Структура:** пункт **1** — цель и база; пункт **2** — разработка (2.1–2.5 текстом, **2.6** — иллюстрация при наличии файла и **листинги кода**); сводная таблица; заключение.

{intro_note}

## 1 Цель, задачи, база практики

### 1.1 Цель и задачи

Целью производственной практики является применение теоретических знаний, полученных при изучении профессионального модуля ПМ 09 «Проектирование, разработка и оптимизация веб-приложений», при разработке программного продукта POTOK. Практика направлена на формирование устойчивых навыков проектирования интерфейсов, реализации пользовательских сценариев, интеграции с серверными сервисами и подготовки сопроводительной документации с пояснениями и листингами.

### 1.2 База практики

Практика выполнялась в рамках учебного проекта в репозитории potok. Технологический стек: Next.js, TypeScript, Supabase, Pusher, Plyr. Работа организована с использованием системы контроля версий, переменных окружения для конфигурации и проверки ключевых сценариев после изменений.

## 2 Разработка веб-приложения

Нормативное задание для сдачи оформлено как **docs/tehnicheskoe_zadanie_potok.pdf** (исходная разметка — TECH_SPEC.md; при необходимости PDF собирается скриптом `scripts/generate_tz_pdf.py`). Макет интерфейса в Figma служил ориентиром по компоновке экранов; далее макет был перенесён в кодовую базу Next.js (App Router) с разбиением на страницы и компоненты.

### 2.1 Техническое задание и постановка задачи

Реализация ведётся в соответствии с требованиями ТЗ к MVP: роли пользователей, видео, категории, комментарии, уведомления, режимы real-time, нефункциональные требования к модульности и разделению клиентской и серверной логики. Архитектурно приложение опирается на связку Next.js и Supabase. К актуальной редакции ТЗ добавлены положения о **модерации**: роли `moderator` и `admin`, очередь жалоб (`reports`), возможность пользователя пожаловаться на видео, комментарий или канал; **эвристические рекомендации** на главной ленте и в боковой панели просмотра (просмотры, лайки, свежесть, подписки, история); **ответы на комментарии** в один уровень вложенности, уведомление адресату об ответе, отметки автора канала («сердце автора») на комментариях без ограничения их числа. Нормативный документ в оформленном виде: **docs/tehnicheskoe_zadanie_potok.pdf** (исходник разметки — `TECH_SPEC.md`, сборка PDF скриптом `scripts/generate_tz_pdf.py`).

### 2.2 Интерфейс: каркас страниц, маршруты и адаптивная вёрстка (подробное изложение)

{section_2_2_long_narrative()}

### 2.3 База данных

Схема данных и политики безопасности на уровне строк (RLS) задаются SQL-скриптами в каталоге **web/supabase**. В миграции **15_moderation_recommendations.sql** добавлены таблицы `reports` (жалобы на видео, комментарий или канал) и `comment_author_hearts` (отметки автора канала на комментариях), поля `role`, `banned_until` и `ban_reason_code` в `public.users`; триггер ограничивает ответы на комментарии **одним уровнем**; при ответе создаётся запись в `notifications`. Политики RLS дополнены: модератор может удалять комментарии и изменять видимость видео; пользователь с активным баном не может создавать комментарии. На стороне приложения используются клиенты Supabase (`web/src/lib/supabase/client.ts`, `server.ts`, `service.ts`). Детальные листинги SQL приведены в п. 2.6.

### 2.4 Авторизация и восстановление доступа

Используется Supabase Auth: регистрация и вход по электронной почте и паролю, сценарии восстановления доступа с подтверждением через письмо. Страницы расположены в `web/src/app/auth/` и связанных маршрутах; шаблоны писем — в `web/supabase/email-templates/`. Роль пользователя в приложении (`user`, `moderator`, `admin`) хранится в таблице `public.users` и проверяется на сервере при доступе к маршрутам `/api/moderation/*` и к странице **/admin**.

### 2.5 Серверная логика, данные и сопоставление с аналогами

Данные запрашиваются через Supabase SDK; медиафайлы размещаются в Storage; для realtime используется Pusher и серверный маршрут `web/src/app/api/realtime/pusher/trigger/route.ts`. Добавлены HTTP-маршруты **POST /api/reports** (подача жалобы), **GET /api/moderation/reports** и **PATCH /api/moderation/reports/[id]** (очередь модератора, разбор жалобы, при необходимости **бан** пользователя через сервисный ключ Supabase). Рекомендации для ленты и боковой панели «Следующее» на странице просмотра считаются в модуле `web/src/lib/recommendations.ts`. Доступ к таблицам регулируется политиками RLS. В качестве ориентира по функциональности выступают массовые видеоплатформы; полный их функционал в учебном MVP не копировался — реализован объём, определённый ТЗ.

### 2.6 Иллюстрации и листинги исходного кода

{shot_md}{listings}

### Сводная таблица реализованных модулей

| Модуль | Статус | Комментарий |
|---|---|---|
| Header + Sidebar | реализовано | каркас навигации, адаптив |
| Главная + поиск | реализовано | лента, категории, поиск, эвристика рекомендаций |
| Просмотр видео | реализовано | плеер, комментарии, ответы, жалобы, сердце автора |
| Канал + студия | реализовано | публикация, контент |
| Авторизация | реализовано | вход, регистрация, OTP |
| Reset Password | реализовано | многошаговый сценарий |
| Supabase + Pusher | реализовано | данные, Storage, realtime |
| Модерация | реализовано | жалобы, роли, панель /admin, API |

## Заключение

В ходе практики реализован рабочий MVP платформы POTOK в соответствии с техническим заданием. В пп. 2.1–2.5 изложены постановка задачи (включая модерацию и рекомендации), интерфейс (п. 2.2), база данных и RLS (п. 2.3), авторизация и роли (п. 2.4), серверная логика и API жалоб (п. 2.5). Подробное повествование о работе над интерфейсом и адаптивом приведено в п. 2.2; иллюстрация (при наличии файла) и листинги исходного кода — в п. 2.6.

Исполнитель __________________ дата __________
"""


def _resolve_image_path(rel: str) -> Path:
    p = Path(rel)
    if p.is_absolute():
        return p
    return (ROOT / rel).resolve()


def _picture_path_for_docx(img_abs: Path) -> Path | None:
    """
    python-docx не принимает WebP и некоторые «переименованные» PNG.
    При ошибке чтения конвертируем во временный настоящий PNG через Pillow.
    """
    try:
        from PIL import Image  # type: ignore[import-untyped]

        im = Image.open(img_abs)
        fmt = (getattr(im, "format", None) or "").upper()
        if fmt in ("WEBP", "MPO") or img_abs.suffix.lower() != ".png":
            tmp = img_abs.parent / f".{img_abs.stem}_docx_embed.png"
            rgb = im.convert("RGB")
            rgb.save(tmp, "PNG", optimize=True)
            return tmp
    except ImportError:
        pass
    except Exception:
        pass
    return img_abs


def _add_picture_safe(doc: Document, img_abs: Path, width_inches: float) -> bool:
    path_try = _picture_path_for_docx(img_abs)
    candidates = [p for p in {img_abs, path_try} if p is not None]
    for cand in candidates:
        if not cand.exists():
            continue
        try:
            doc.add_picture(str(cand), width=Inches(width_inches))
            if cand.name.startswith(".") and "_docx_embed.png" in cand.name:
                try:
                    cand.unlink(missing_ok=True)
                except OSError:
                    pass
            return True
        except (UnrecognizedImageError, OSError):
            continue
    return False


def _add_code_listing_paragraph(doc: Document, block: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_together = True
    run = p.add_run(block)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def _add_subheading_plain(doc: Document, text: str) -> None:
    """Подзаголовок без уровня в нумерации Word (не создаёт 2.2.1–2.2.7 для #### в Markdown)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def markdown_to_docx_verbose(md_text: str, output_path: Path) -> None:
    """Как markdown_to_docx в practice_report, плюс картинки ![alt](path)."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    in_code = False
    code_lines: list[str] = []
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        raw_line = lines[i]
        line = raw_line.rstrip("\n")
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                # Закрытие блока кода: один абзац-листинг моноширинным шрифтом
                block = "\n".join(code_lines)
                code_lines = []
                in_code = False
                _add_code_listing_paragraph(doc, block)
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not stripped:
            doc.add_paragraph("")
            i += 1
            continue

        m_img = IMG_LINE.match(stripped)
        if m_img:
            alt, rel = m_img.groups()
            img_abs = _resolve_image_path(rel)
            if img_abs.exists():
                if alt:
                    doc.add_paragraph(alt)
                if not _add_picture_safe(doc, img_abs, 6.2):
                    doc.add_paragraph(
                        f"[Не удалось вставить изображение: {rel}. "
                        "Сохраните файл как PNG или JPEG и установите Pillow: pip install pillow]"
                    )
            else:
                doc.add_paragraph(
                    f"[Рисунок не найден по пути {rel} — поместите файл и пересоберите отчёт. Подпись: {alt}]"
                )
            i += 1
            continue

        if stripped.startswith("# "):
            h = doc.add_heading(stripped[2:].strip(), level=0)
            h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
            i += 1
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
            i += 1
            continue
        if stripped.startswith("#### "):
            _add_subheading_plain(doc, stripped[5:].strip())
            i += 1
            continue
        if stripped.startswith("- "):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
            i += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            table_block: list[str] = []
            while i < len(lines):
                candidate = lines[i].rstrip("\n")
                if candidate.strip().startswith("|") and candidate.strip().endswith("|"):
                    table_block.append(candidate)
                    i += 1
                else:
                    break
            if len(table_block) >= 2:
                headers = [c.strip() for c in table_block[0].strip("|").split("|")]
                sep = table_block[1].replace("-", "").replace("|", "").replace(":", "").strip()
                data_rows = table_block[2:] if sep == "" else table_block[1:]
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = "Table Grid"
                for col, value in enumerate(headers):
                    table.rows[0].cells[col].text = value
                for row_line in data_rows:
                    values = [c.strip() for c in row_line.strip("|").split("|")]
                    row = table.add_row().cells
                    for col in range(min(len(values), len(headers))):
                        row[col].text = values[col]
                doc.add_paragraph("")
                continue
            doc.add_paragraph(line)
            i += 1
            continue

        # Курсив *...* упрощаем до обычного абзаца
        if stripped.startswith("*") and stripped.endswith("*") and len(stripped) > 2:
            doc.add_paragraph(stripped[1:-1].strip())
            i += 1
            continue

        doc.add_paragraph(line)
        i += 1

    if in_code and code_lines:
        _add_code_listing_paragraph(doc, "\n".join(code_lines))

    doc.save(output_path)


def markdown_to_html_verbose(md_text: str) -> str:
    """HTML с поддержкой изображений и кода."""
    lines = md_text.splitlines()
    html: list[str] = [
        "<!doctype html>",
        "<html lang='ru'>",
        "<head>",
        "  <meta charset='utf-8' />",
        "  <meta name='viewport' content='width=device-width, initial-scale=1' />",
        "  <title>Отчёт по практике — POTOK (подробно)</title>",
        "  <style>",
        "    body{font-family:'Times New Roman',Georgia,serif;max-width:920px;margin:28px auto;line-height:1.55;color:#1a1a1a;}",
        "    h1{text-align:center;font-size:1.35rem;margin-bottom:1.2em;}",
        "    h2{font-size:1.2rem;margin:1.4em 0 0.6em;border-bottom:1px solid #ccc;padding-bottom:0.2em;}",
        "    h3,h4{margin:1em 0 0.5em;}",
        "    p{margin:0.55em 0;text-align:justify;}",
        "    pre{background:#f6f7f9;padding:12px;overflow:auto;border:1px solid #dde1e6;border-radius:4px;font-size:0.88rem;}",
        "    code{font-family:Consolas,'Courier New',monospace;}",
        "    figure{margin:1.2em 0;text-align:center;}",
        "    figure img{max-width:100%;height:auto;border:1px solid #ccc;border-radius:4px;box-shadow:0 4px 16px rgba(0,0,0,0.08);}",
        "    figcaption{font-size:0.92rem;color:#444;margin-top:0.5em;font-style:italic;}",
        "    table{border-collapse:collapse;width:100%;margin:12px 0;font-size:0.95rem;}",
        "    th,td{border:1px solid #555;padding:8px 10px;vertical-align:top;}",
        "    th{background:#f0f0f0;}",
        "    ul{margin:8px 0 8px 24px;}",
        "  </style>",
        "</head>",
        "<body>",
    ]

    in_code = False
    i = 0
    while i < len(lines):
        line = lines[i].rstrip("\n")
        stripped = line.strip()

        if stripped.startswith("```"):
            if not in_code:
                html.append("<pre><code>")
                in_code = True
            else:
                html.append("</code></pre>")
                in_code = False
            i += 1
            continue

        if in_code:
            escaped = (
                line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            )
            html.append(escaped)
            i += 1
            continue

        m_img = IMG_LINE.match(stripped)
        if m_img:
            alt, rel = m_img.groups()
            rel_clean = rel.replace("\\", "/")
            if rel_clean.startswith("docs/"):
                src = rel_clean[5:]
            else:
                src = rel_clean
            abs_img = _resolve_image_path(rel)
            if abs_img.exists():
                html.append("<figure>")
                html.append(f"<img src='{src}' alt='{alt.replace(chr(39), '&#39;')}' />")
                html.append(f"<figcaption>{alt}</figcaption>")
                html.append("</figure>")
            else:
                html.append(
                    f"<p class='missing'><em>Рисунок недоступен (файл не найден): {rel}</em></p>"
                )
            i += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            table_block: list[str] = []
            while i < len(lines):
                candidate = lines[i].rstrip("\n").strip()
                if candidate.startswith("|") and candidate.endswith("|"):
                    table_block.append(candidate)
                    i += 1
                else:
                    break
            if len(table_block) >= 2:
                header = [c.strip() for c in table_block[0].strip("|").split("|")]
                sep = table_block[1].replace("-", "").replace("|", "").replace(":", "").strip()
                data_rows = table_block[2:] if sep == "" else table_block[1:]
                html.append("<table><thead><tr>")
                for h in header:
                    html.append(f"<th>{h}</th>")
                html.append("</tr></thead><tbody>")
                for row in data_rows:
                    html.append("<tr>")
                    for cell in [c.strip() for c in row.strip("|").split("|")]:
                        html.append(f"<td>{cell}</td>")
                    html.append("</tr>")
                html.append("</tbody></table>")
                continue

        if not stripped:
            html.append("<p></p>")
            i += 1
            continue
        if stripped.startswith("# "):
            html.append(f"<h1>{stripped[2:]}</h1>")
            i += 1
            continue
        if stripped.startswith("## "):
            html.append(f"<h2>{stripped[3:]}</h2>")
            i += 1
            continue
        if stripped.startswith("### "):
            html.append(f"<h3>{stripped[4:]}</h3>")
            i += 1
            continue
        if stripped.startswith("#### "):
            html.append(f"<h4>{stripped[5:]}</h4>")
            i += 1
            continue
        if stripped.startswith("- "):
            html.append(f"<ul><li>{stripped[2:]}</li></ul>")
            i += 1
            continue
        if stripped.startswith("*") and stripped.endswith("*") and len(stripped) > 2:
            html.append(f"<p><em>{stripped[1:-1]}</em></p>")
            i += 1
            continue

        # экранирование HTML в обычных абзацах
        esc = (
            stripped.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )
        html.append(f"<p>{esc}</p>")
        i += 1

    html.extend(["</body>", "</html>"])
    return "\n".join(html)


def ensure_screenshots_dir() -> None:
    d = ROOT / "docs" / "screenshots"
    d.mkdir(parents=True, exist_ok=True)
    readme = d / "README.txt"
    if not readme.exists():
        readme.write_text(
            "Положите сюда файл otchet_razdel_2_2.png — скриншот фрагмента отчёта в Word\n"
            "(раздел об интерфейсе), чтобы он подставился в подробный отчёт.\n"
            "Скрипт: scripts/practice_report_verbose.py\n",
            encoding="utf-8",
        )


def try_copy_default_screenshot(dest_rel: str) -> None:
    """Если в типичном месте Cursor лежит скрин — копируем в docs/screenshots."""
    dest = ROOT / dest_rel
    if dest.exists():
        return
    candidates = [
        ROOT / "assets" / "otchet_razdel_2_2.png",
        Path(
            os.environ.get("USERPROFILE", "")
        )
        / ".cursor"
        / "projects"
        / "c-Users-Danie-OneDrive-Desktop-potok"
        / "assets",
    ]
    for folder in candidates:
        if folder.is_dir():
            for png in folder.glob("*.png"):
                shutil.copy2(png, dest)
                return
        elif folder.suffix.lower() == ".png" and folder.exists():
            shutil.copy2(folder, dest)
            return


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Подробный отчёт POTOK: текст → скриншот → листинги"
    )
    parser.add_argument(
        "--screenshot",
        default=DEFAULT_SCREENSHOT_REL,
        help=f"путь к PNG относительно корня репо (по умолчанию: {DEFAULT_SCREENSHOT_REL})",
    )
    parser.add_argument(
        "--all",
        action="store_true",
        help="дополнительно записать .md, .html и при возможности .pdf (по умолчанию только .docx)",
    )
    args = parser.parse_args()
    os.chdir(ROOT)

    ensure_docs_dir()
    ensure_screenshots_dir()
    try_copy_default_screenshot(args.screenshot)

    body = build_verbose_report_body(args.screenshot)

    if args.all:
        OUT_MD.write_text(body, encoding="utf-8")
        print(f"Отчёт MD:   {OUT_MD.relative_to(ROOT)}")
        html_content = markdown_to_html_verbose(body)
        OUT_HTML.write_text(html_content, encoding="utf-8")

    try:
        markdown_to_docx_verbose(body, OUT_DOCX)
        print(f"Отчёт DOCX: {OUT_DOCX.relative_to(ROOT)}")
    except PermissionError:
        alt = ROOT / "docs" / "otchet_praktika_potok_new.docx"
        markdown_to_docx_verbose(body, alt)
        print(f"Отчёт DOCX (резерв): {alt.relative_to(ROOT)}")

    if args.all:
        if try_export_pdf_from_html(OUT_HTML, OUT_PDF):
            print(f"Отчёт PDF:  {OUT_PDF.relative_to(ROOT)}")
        else:
            print(
                "PDF не сгенерирован автоматически. Откройте HTML в браузере: Печать → Сохранить как PDF."
            )


if __name__ == "__main__":
    main()
